import magic # python-magic
import io
from typing import Callable, Dict, Optional

# Import parsers for specific file types
try:
    import pypdfium2 as pdfium
except ImportError:
    pdfium = None # Handle missing dependency gracefully

try:
    from docx import Document # python-docx
except ImportError:
    Document = None # Handle missing dependency gracefully

from src.config import settings

MAX_TEXT_LENGTH = 1_000_000 # Limit extracted text length to prevent memory issues with huge files

def parse_pdf(file_bytes: bytes) -> str:
    """Extracts text from a PDF file's bytes."""
    if not pdfium:
        raise ImportError("pypdfium2 is not installed. Cannot parse PDF files.")
    try:
        pdf = pdfium.PdfDocument(file_bytes)
        text = ""
        for page in pdf:
            text += page.get_textpage().get_text_range() + "\n"
            if len(text) > MAX_TEXT_LENGTH:
                break # Stop if text becomes too long
        return text[:MAX_TEXT_LENGTH]
    except Exception as e:
        # Log error e
        raise ValueError(f"Failed to parse PDF: {e}") from e

def parse_docx(file_bytes: bytes) -> str:
    """Extracts text from a DOCX file's bytes."""
    if not Document:
        raise ImportError("python-docx is not installed. Cannot parse DOCX files.")
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text[:MAX_TEXT_LENGTH]
    except Exception as e:
        # Log error e
        raise ValueError(f"Failed to parse DOCX: {e}") from e

def parse_txt(file_bytes: bytes) -> str:
    """Extracts text from a TXT file's bytes, attempting common encodings."""
    encodings_to_try = ['utf-8', 'latin-1', 'windows-1252']
    for encoding in encodings_to_try:
        try:
            text = file_bytes.decode(encoding)
            return text[:MAX_TEXT_LENGTH]
        except UnicodeDecodeError:
            continue
    raise ValueError("Failed to decode TXT file with common encodings.")

FILE_PARSERS: Dict[str, Callable[[bytes], str]] = {
    "parse_pdf": parse_pdf,
    "parse_docx": parse_docx,
    "parse_txt": parse_txt,
}

def get_file_mime_type(file_bytes: bytes) -> Optional[str]:
    """Detects the MIME type of a file using python-magic."""
    try:
        mime_type = magic.from_buffer(file_bytes, mime=True)
        return mime_type
    except Exception as e:
        # Log error e
        print(f"Error detecting MIME type: {e}") # Basic logging
        return None

def extract_text_from_file(file_bytes: bytes, original_filename: Optional[str] = None) -> Optional[str]:
    """Detects file type and extracts text content."""
    mime_type = get_file_mime_type(file_bytes)
    
    if not mime_type:
        # Fallback: try to guess from extension if mime_type detection fails or is too generic
        if original_filename and original_filename.lower().endswith(".pdf"):
            mime_type = "application/pdf"
        elif original_filename and (original_filename.lower().endswith(".docx") ):
             mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif original_filename and original_filename.lower().endswith(".txt"):
            mime_type = "text/plain"
        else:
            # Potentially log a warning here if a known extension is not matched
            pass # Allow to proceed if mime_type is still None, parser_name will be None

    parser_name = settings.SUPPORTED_FILE_TYPES.get(mime_type)
    
    if not parser_name:
        # If still no parser and it's a text-like generic mime type, try plain text parsing
        if mime_type and mime_type.startswith("text/"):
            parser_name = "parse_txt"
        else:
            # Log: print(f"Unsupported file type: {mime_type} for file {original_filename}")
            raise ValueError(f"Unsupported or unrecognized file type: {mime_type}")

    parser_func = FILE_PARSERS.get(parser_name)

    if not parser_func:
        # This case should ideally be caught by the check above
        raise ValueError(f"No parser available for determined type linked to parser name: {parser_name}")

    return parser_func(file_bytes) 